import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Set page configuration
st.set_page_config(
    page_title="Drug Sales Data Visualization",
    page_icon="💊",
    layout="wide"
)

# Main title
st.title("Drug Sales Data Visualization Tool")
st.markdown("Upload your drug sales data file (CSV or Excel) to generate visualizations and reports.")

# File upload section
st.header("1. Upload Data")
uploaded_file = st.file_uploader("Choose a CSV or Excel file", type=["csv", "xlsx", "xls"])

# Initialize dataframe
df = None

# Function to create time series plot
def create_time_series_plot(df, selected_drugs):
    fig, ax = plt.subplots(figsize=(12, 6))
    
    for drug in selected_drugs:
        ax.plot(df.index, df[drug], label=drug)
    
    ax.set_xlabel('Date')
    ax.set_ylabel('Sales Volume')
    ax.set_title('Monthly Drug Sales Volume Over Time')
    ax.legend()
    ax.grid(True, alpha=0.3)
    
    plt.tight_layout()
    return fig

# Function to create boxplot
def create_boxplot(df, selected_drugs):
    fig, ax = plt.subplots(figsize=(12, 6))
    
    # Reshape data for boxplot
    plot_data = df[selected_drugs].melt(var_name='Drug Category', value_name='Sales Volume')
    
    # Create boxplot
    sns.boxplot(x='Drug Category', y='Sales Volume', data=plot_data, ax=ax)
    
    ax.set_xlabel('Drug Category')
    ax.set_ylabel('Sales Volume')
    ax.set_title('Distribution of Drug Sales Volume by Category')
    plt.xticks(rotation=45)
    
    plt.tight_layout()
    return fig

# Function to create heatmap
def create_heatmap(df, selected_drugs):
    # Calculate correlation matrix
    corr_matrix = df[selected_drugs].corr()
    
    # Create heatmap
    fig, ax = plt.subplots(figsize=(10, 8))
    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', ax=ax, fmt=".2f", linewidths=0.5)
    ax.set_title('Correlation Between Drug Sales Categories')
    
    plt.tight_layout()
    return fig

# Function to create bar chart of annual sales
def create_annual_sales(df, selected_drugs):
    # Create a copy of the dataframe with a year column
    df_yearly = df.copy()
    df_yearly['Year'] = df_yearly.index.year
    
    # Group by year and sum
    yearly_sales = df_yearly.groupby('Year')[selected_drugs].sum()
    
    # Create bar chart
    fig, ax = plt.subplots(figsize=(12, 6))
    yearly_sales[selected_drugs].plot(kind='bar', ax=ax)
    
    ax.set_xlabel('Year')
    ax.set_ylabel('Total Sales Volume')
    ax.set_title('Annual Drug Sales by Category')
    ax.legend(title='Drug Category')
    
    plt.tight_layout()
    return fig

# Function to create a PDF report
def create_pdf_report(fig, description):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Add title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, height - 72, "Drug Sales Data Visualization Report")
    
    # Add description
    c.setFont("Helvetica", 12)
    c.drawString(72, height - 100, "Description:")
    
    # Wrap text for description
    text_object = c.beginText(72, height - 120)
    text_object.setFont("Helvetica", 10)
    
    # Split description into lines
    words = description.split()
    lines = []
    line = ""
    for word in words:
        if len(line + word) + 1 <= 80:  # Limit line length
            line += word + " "
        else:
            lines.append(line)
            line = word + " "
    if line:
        lines.append(line)
    
    for line in lines:
        text_object.textLine(line)
    
    c.drawText(text_object)
    
    # Save figure to temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        temp_filename = tmp_file.name
        fig.savefig(temp_filename, format='png', dpi=300, bbox_inches='tight')
    
    # Add figure to PDF
    c.drawImage(temp_filename, 72, 100, width=width-144, height=height-300)
    
    c.save()
    os.unlink(temp_filename)  # Remove temporary file
    
    buffer.seek(0)
    return buffer

# Function to create a Word document report
def create_word_report(fig, description):
    doc = Document()
    doc.add_heading('Drug Sales Data Visualization Report', 0)
    
    doc.add_heading('Description:', level=1)
    doc.add_paragraph(description)
    
    # Save figure to temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        temp_filename = tmp_file.name
        fig.savefig(temp_filename, format='png', dpi=300, bbox_inches='tight')
    
    # Add figure to Word document
    doc.add_heading('Visualization:', level=1)
    doc.add_picture(temp_filename, width=Inches(6))
    
    # Save the document to a BytesIO object
    buffer = io.BytesIO()
    doc.save(buffer)
    os.unlink(temp_filename)  # Remove temporary file
    
    buffer.seek(0)
    return buffer

# Function to create download link
def get_download_link(buffer, download_filename, link_text):
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{download_filename}">{link_text}</a>'

# Process uploaded file
if uploaded_file is not None:
    try:
        # Check file type
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        else:  # Excel file
            df = pd.read_excel(uploaded_file)
        
        # Display raw data sample
        st.subheader("Data Preview")
        st.dataframe(df.head())
        
        # Check if 'datum' column exists (date column)
        date_column = None
        if 'datum' in df.columns:
            date_column = 'datum'
        elif 'date' in df.columns:
            date_column = 'date'
        
        # If date column exists, set it as index
        if date_column:
            # Convert to datetime
            df[date_column] = pd.to_datetime(df[date_column])
            df.set_index(date_column, inplace=True)
        
        # Get drug categories (all columns except date column)
        if date_column:
            drug_categories = list(df.columns)
        else:
            # If no date column found, assume first column might be date
            st.warning("No explicit date column found. Assuming first column is for dates.")
            date_column = df.columns[0]
            df[date_column] = pd.to_datetime(df[date_column])
            df.set_index(date_column, inplace=True)
            drug_categories = list(df.columns)
        
        # Visualization options section
        st.header("2. Create Visualization")
        
        # Select drugs for visualization
        selected_drugs = st.multiselect(
            "Select drug categories to visualize:",
            options=drug_categories,
            default=drug_categories[:3] if len(drug_categories) >= 3 else drug_categories
        )
        
        # Chart type selection
        chart_type = st.selectbox(
            "Select chart type:",
            options=["Time Series", "Box Plot", "Correlation Heatmap", "Annual Sales"]
        )
        
        # Graph title and description
        graph_title = st.text_input("Graph Title", "Drug Sales Analysis")
        
        graph_description = st.text_area(
            "Graph Description (max 80 words):",
            "This visualization shows the trends in drug sales volume over the analyzed period. "
            "The data reveals seasonal patterns and overall market trends for various pharmaceutical categories.",
            max_chars=400
        )
        
        # Create visualization button
        if st.button("Generate Visualization"):
            if not selected_drugs:
                st.error("Please select at least one drug category.")
            else:
                st.subheader("3. Generated Visualization")
                
                # Create appropriate chart based on selection
                if chart_type == "Time Series":
                    fig = create_time_series_plot(df, selected_drugs)
                elif chart_type == "Box Plot":
                    fig = create_boxplot(df, selected_drugs)
                elif chart_type == "Correlation Heatmap":
                    fig = create_heatmap(df, selected_drugs)
                elif chart_type == "Annual Sales":
                    fig = create_annual_sales(df, selected_drugs)
                
                # Display chart
                st.pyplot(fig)
                
                # Download options
                st.header("4. Download Report")
                col1, col2 = st.columns(2)
                
                # Create PDF buffer
                pdf_buffer = create_pdf_report(fig, graph_description)
                pdf_link = get_download_link(pdf_buffer, "drug_sales_report.pdf", "Download PDF Report")
                
                # Create Word buffer
                docx_buffer = create_word_report(fig, graph_description)
                docx_link = get_download_link(docx_buffer, "drug_sales_report.docx", "Download Word Report")
                
                with col1:
                    st.markdown(pdf_link, unsafe_allow_html=True)
                
                with col2:
                    st.markdown(docx_link, unsafe_allow_html=True)
                
    except Exception as e:
        st.error(f"Error processing file: {e}")
else:
    st.info("Please upload a file to begin analysis.")

# Footer
st.markdown("---")
st.markdown("### Instructions")
st.markdown("""
1. Upload your drug sales data file in CSV or Excel format
2. Select the drug categories you want to visualize
3. Choose a chart type and add a title and description
4. Generate the visualization
5. Download the report in PDF or Word format
""")